import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, Toplevel
from docx import Document
import os

plantilla_paths = []
directorio_guardado = ""
clausulas = {}
clausulas_seleccionadas = []
boceto_contrato = ""

def cargar_clausulas():
    carpeta_clausulas = "clausulas"
    if not os.path.exists(carpeta_clausulas):
        messagebox.showwarning("Advertencia", "No se encontró la carpeta de cláusulas.")
        return

    for nombre_archivo in os.listdir(carpeta_clausulas):
        ruta_archivo = os.path.join(carpeta_clausulas, nombre_archivo)
        if os.path.isfile(ruta_archivo):
            nombre_clausula = os.path.splitext(nombre_archivo)[0]
            try:
                if nombre_archivo.endswith(".txt"):
                    with open(ruta_archivo, "r", encoding="utf-8") as archivo:
                        clausulas[nombre_clausula] = archivo.read()
                elif nombre_archivo.endswith(".docx"):
                    document = Document(ruta_archivo)
                    contenido_documento = ""
                    for parrafo in document.paragraphs:
                        contenido_documento += parrafo.text + "\n"
                    clausulas[nombre_clausula] = contenido_documento
            except Exception as e:
                messagebox.showerror("Error", f"Error al leer {nombre_archivo}: {e}")

def mostrar_vista_previa(event):
    seleccion = listbox_clausulas.curselection()
    if seleccion:
        nombre_clausula = listbox_clausulas.get(seleccion[0])
        texto_clausula = clausulas[nombre_clausula]

        ventana_vista_previa = Toplevel(root)
        ventana_vista_previa.title(f"Vista previa de {nombre_clausula}")

        texto_vista_previa = scrolledtext.ScrolledText(ventana_vista_previa, width=60, height=20)
        texto_vista_previa.insert(tk.END, texto_clausula)
        texto_vista_previa.config(state=tk.DISABLED)
        texto_vista_previa.pack(padx=10, pady=10)

def actualizar_boceto():
    global boceto_contrato

    boceto_contrato = ""
    if plantilla_paths:
        for plantilla in plantilla_paths:
            doc = Document(plantilla)
            for p in doc.paragraphs:
                boceto_contrato += p.text + "\n"

    texto_boceto.config(state=tk.NORMAL)
    texto_boceto.delete("1.0", tk.END)
    texto_boceto.insert(tk.END, boceto_contrato)

    for nombre_clausula in clausulas_seleccionadas:
        texto_boceto.insert(tk.END, "\n\n" + clausulas[nombre_clausula])

    texto_boceto.config(state=tk.DISABLED)

def seleccionar_clausula(event):
    global clausulas_seleccionadas

    seleccion = listbox_clausulas.curselection()
    clausulas_seleccionadas = [listbox_clausulas.get(i) for i in seleccion]
    actualizar_boceto()

def seleccionar_directorio():
    global directorio_guardado
    directorio_guardado = filedialog.askdirectory()
    if directorio_guardado:
        lbl_directorio.config(text=f"Carpeta de guardado: {directorio_guardado}")

def generar_contrato():
    if not directorio_guardado:
        messagebox.showwarning("Advertencia", "Por favor, selecciona una carpeta para guardar los archivos.")
        return

    datos = {
        "Nombre": entry_NOMBREPFISICA.get(),
        "numeroID": entry_NUMEROID.get(),
        "institucion": entry_institucion.get(),
        "curp": entry_curp.get(),
        "rfc": entry_rfc.get(),
        "direccion": entry_direccion.get("1.0", tk.END).strip(),
        "correo": entry_correo.get(),
        "FECHAF": entry_FECHAF.get(),
    }

    if not all(datos.values()):
        messagebox.showwarning("Advertencia", "Por favor, completa todos los campos.")
        return

    nombre_archivo_base = entry_nombre_archivo.get().strip()
    if not nombre_archivo_base:
        messagebox.showwarning("Advertencia", "Por favor, ingresa un nombre de archivo.")
        return

    nombre_archivo_base = f"{datos['NombrePFISICA']}_{nombre_archivo_base}.docx"

    doc = Document()

    doc.add_paragraph(f"Nombre: {datos['NombrePFISICA']}")
    doc.add_paragraph(f"Número de Identificación: {datos['numeroID']}")
    doc.add_paragraph(f"Institución: {datos['institucion']}")
    doc.add_paragraph(f"CURP: {datos['curp']}")
    doc.add_paragraph(f"RFC: {datos['rfc']}")
    doc.add_paragraph(f"Dirección: {datos['direccion']}")
    doc.add_paragraph(f"Correo: {datos['correo']}")
    doc.add_paragraph(f"Fecha: {datos['FECHAF']}")

    for nombre_clausula in clausulas_seleccionadas:
        doc.add_paragraph(clausulas[nombre_clausula])

    ruta_completa = os.path.join(directorio_guardado, nombre_archivo_base)
    doc.save(ruta_completa)

    messagebox.showinfo("Éxito", f"Contrato generado: {nombre_archivo_base}")
    actualizar_boceto()

def abrir_carpeta():
    if directorio_guardado:
        os.startfile(directorio_guardado)
    else:
        messagebox.showwarning("Advertencia", "No hay carpeta seleccionada.")

root = tk.Tk()
root.title("Generador de Contratos Personalizado")
root.geometry("1200x700")

frame_izquierda = tk.Frame(root)
frame_izquierda.pack(side=tk.LEFT, padx=10, pady=10)

frame_derecha = tk.Frame(root)
frame_derecha.pack(side=tk.RIGHT, padx=10, pady=10)

frame_centro = tk.Frame(root)
frame_centro.pack(padx=10, pady=10)

btn_plantilla = tk.Button(frame_izquierda, text="Seleccionar Plantillas", width=20, command=seleccionar_plantillas)
btn_plantilla.pack(pady=5)

lbl_plantilla = tk.Label(frame_izquierda, text="No se ha seleccionado ninguna plantilla", wraplength=200)
lbl_plantilla.pack()

btn_directorio = tk.Button(frame_izquierda, text="Seleccionar Carpeta", width=20, command=seleccionar_directorio)
btn_directorio.pack(pady=5)

lbl_directorio = tk.Label(frame_izquierda, text="No se ha seleccionado carpeta", wraplength=200)
lbl_directorio.pack()

btn_generar = tk.Button(frame_izquierda, text="Generar Contrato", width=20, command=generar_contrato)
btn_generar.pack(pady=5)

btn_abrir = tk.Button(frame_izquierda, text="Abrir Carpeta", width=20, command=abrir_carpeta)
btn_abrir.pack(pady=5)

lbl_clausulas = tk.Label(frame_izquierda, text="Cláusulas disponibles:")
lbl_clausulas.pack()

scrollbar_clausulas = tk.Scrollbar(frame_izquierda)
scrollbar_clausulas.pack(side=tk.RIGHT, fill=tk.Y)

listbox_clausulas = tk.Listbox(frame_izquierda, selectmode=tk.MULTIPLE, yscrollcommand=scrollbar_clausulas.set, width=40, height=10)
listbox_clausulas.pack()

scrollbar_clausulas.config(command=listbox_clausulas.yview)

cargar_clausulas()
for clausula in clausulas:
    listbox_clausulas.insert(tk.END, clausula)

listbox_clausulas.bind("<Double-Button-1>", mostrar_vista_previa)
listbox_clausulas.bind("<<ListboxSelect>>", seleccionar_clausula)

lbl_boceto = tk.Label(frame_centro, text="Boceto del Contrato:")
lbl_boceto.pack()

texto_boceto = scrolledtext.ScrolledText(frame_centro, width=80, height=20)
texto_boceto.pack()
texto_boceto.config(state=tk.DISABLED)

tk.Label(frame_derecha, text="NombrePFISICA:").pack()
entry_NOMBREPFISICA = tk.Entry(frame_derecha, width=30)
entry_NOMBREPFISICA.pack()

tk.Label(frame_derecha, text="NumeroID:").pack()
entry_NUMEROID = tk.Entry(frame_derecha, width=30)
entry_NUMEROID.pack()

tk.Label(frame_derecha, text="Institucion:").pack()
entry_institucion = tk.Entry(frame_derecha, width=30)
entry_institucion.pack()

tk.Label(frame_derecha, text="CURP:").pack()
entry_curp = tk.Entry(frame_derecha, width=30)
entry_curp.pack()

tk.Label(frame_derecha, text="RFC:").pack()
entry_rfc = tk.Entry(frame_derecha, width=30)
entry_rfc.pack()

tk.Label(frame_derecha, text="Direccion:").pack()
entry_direccion = scrolledtext.ScrolledText(frame_derecha, height=4, width=30)
entry_direccion.pack()

tk.Label(frame_derecha, text="Correo:").pack()
entry_correo = tk.Entry(frame_derecha, width=30)
entry_correo.pack()

tk.Label(frame_derecha, text="Fecha:").pack()
entry_FECHAF = tk.Entry(frame_derecha, width=30)
entry_FECHAF.pack()

tk.Label(frame_derecha, text="Nombre del Archivo:").pack()
entry_nombre_archivo = tk.Entry(frame_derecha, width=30)
entry_nombre_archivo.pack()

root.mainloop()