from flask import Flask, request, jsonify
from docx import Document

app = Flask(__name__)

def generar_documento(datos):
    doc = Document()
    doc.add_heading('Documento Legal', level=1)
    
    doc.add_paragraph(f"Actor: {datos.get('actor', 'No especificado')}")
    doc.add_paragraph(f"Demandado: {datos.get('demandado', 'No especificado')}")
    doc.add_paragraph(f"Documento base: {datos.get('documento_base', 'No especificado')}")
    doc.add_paragraph(f"Otras pruebas: {datos.get('otras_pruebas', 'No especificado')}")
    
    filename = "documento_legal.docx"
    doc.save(filename)
    return filename

@app.route('/generar_documento', methods=['POST'])
def recibir_datos():
    datos = request.json
    nombre_archivo = generar_documento(datos)
    return jsonify({"mensaje": "Documento generado con éxito", "archivo": nombre_archivo})

if __name__ == '__main__':
    app.run(debug=True)
